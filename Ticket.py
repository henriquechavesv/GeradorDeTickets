from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import os
from datetime import datetime, timedelta

# Defina o caminho da pasta "Tickets..."
pasta_tickets_cafe = "Tickets Café da manhã"
pasta_tickets_almoco = "Tickets Almoço"
pasta_tickets_janta = "Tickets Janta"

# Função para lidar com a correspondência de benefícios
def match_case(refeicao, pasta_tickets_cafe, pasta_tickets_almoco, pasta_tickets_janta, nome):
    match refeicao:
        case "Café da manhã":
                doc.save(os.path.join(pasta_tickets_cafe, f'Ticket - {nome}.docx'))
        case "Almoço":
                doc.save(os.path.join(pasta_tickets_almoco, f'Ticket - {nome}.docx'))
        case "Janta":
                doc.save(os.path.join(pasta_tickets_janta, f'Ticket - {nome}.docx'))

# Entradas do usuário
refeicao = input('Digite qual refeição será usada nesse ticket: ')
horario = input('Digite plantão ou comercial: ')


if horario == 'Comercial':
    print('Digite o nome do funcionário ou F para finalizar o programa: ') 
    while True: 
        nome = input('')
        if nome == 'F': 
            break
        else:
            # Defina as medidas em centímetros (assinatura)
            largura_cm = 7.00 
            altura_cm = 1.50   

            # Converta as medidas de centímetros para polegadas (assinatura)
            largura_polegadas = largura_cm / 2.54
            altura_polegadas = altura_cm / 2.54

            # Defina as medidas em centímetros da logo
            largura_logo_cm = 9.00 
            altura_logo_cm = 1.50   

            # Converta as medidas de centímetros para polegadas (logo do hospital)
            largura_polegadas_logo = largura_logo_cm / 2.54
            altura_polegadas_logo = altura_logo_cm / 2.54

            # Crie um novo documento
            doc = Document()

            # Adicione uma tabela 6x5 ao documento
            table = doc.add_table(rows=6, cols=4)
            table.autofit = False

            # Configure as margens para zero e margem da esquerda para 10
            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(0)
                section.bottom_margin = Mm(0)
                section.left_margin = Mm(10)
                section.right_margin = Mm(0)

            # Crie uma lista para armazenar os dias da semana
            data = '   /    /    '
            
            # Configure o espaçamento entre as células
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.width = Cm(5)  # Largura da célula (ajuste conforme necessário)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(60)  # Espaço após o parágrafo
                    cell.paragraphs[0].paragraph_format.space_before = Pt(18)  # Espaço antes do parágrafo

            # Preencha a tabela com as informações nas variáveis
            for idx, row in enumerate(table.rows):
                for cell in row.cells:
                    # Crie um novo parágrafo para a célula
                    paragraph = cell.paragraphs[0]

                    # Antes de adicionar o texto, adicione a imagem
                    imagem_logo = "Imagens/logo.jpg"  # Substitua pelo caminho da sua imagem
                    run = paragraph.add_run()
                    run.add_picture(imagem_logo, width=Cm(largura_polegadas_logo), height=Cm(altura_polegadas_logo))  # Ajuste as dimensões da imagem

                    # Alinhe o parágrafo ao meio
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinhamento ao centro
                    lista_nomes = []
                    lista_nomes.append(nome)
                    # Adicione o texto formatado
                    paragraph.add_run(f"Nome: {nome}\nData: {data}\nBenefício: {refeicao}")
                    
                    paragraph = cell.paragraphs[0]

                    # Adicione a imagem da assinatura digital
                    assinatura_path = "Imagens/assinatura.png"  # Substitua pelo caminho correto da imagem
                    paragraph.add_run().add_picture(assinatura_path, width=Cm(largura_polegadas), height=Cm(altura_polegadas))

            # Chame a função match_case para salvar o documento com base no benefício
            match_case(refeicao, pasta_tickets_cafe, pasta_tickets_almoco, pasta_tickets_janta, nome)


elif horario == 'Plantão':
    data = input('Digite a data (dd/mm/aaaa): ')
    print('Digite o nome do funcionário ou F para finalizar o programa: ')
    while True: 
        nome = input('')
        if nome == 'F': 
            break
        else:
            # Defina as medidas em centímetros (assinatura)
            largura_cm = 7.00 
            altura_cm = 1.50   

            # Converta as medidas de centímetros para polegadas (assinatura)
            largura_polegadas = largura_cm / 2.54
            altura_polegadas = altura_cm / 2.54

            # Defina as medidas em centímetros da logo
            largura_logo_cm = 9.00 
            altura_logo_cm = 1.50   

            # Converta as medidas de centímetros para polegadas (logo do hospital)
            largura_polegadas_logo = largura_logo_cm / 2.54
            altura_polegadas_logo = altura_logo_cm / 2.54

            # Crie um novo documento
            doc = Document()

            # Adicione uma tabela 6x5 ao documento
            table = doc.add_table(rows=4, cols=4)
            table.autofit = False

            # Configure as margens para zero e margem da esquerda para 10
            sections = doc.sections
            for section in sections:
                section.top_margin = Mm(0)
                section.bottom_margin = Mm(0)
                section.left_margin = Mm(10)
                section.right_margin = Mm(0)

            # Crie uma lista para armazenar as datas incrementadas
            datas_incrementadas = [data]

            # Configure o espaçamento entre as células
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.width = Cm(5)  # Largura da célula (ajuste conforme necessário)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(60)  # Espaço após o parágrafo
                    cell.paragraphs[0].paragraph_format.space_before = Pt(20)  # Espaço antes do parágrafo

            # Preencha a tabela com as informações nas variáveis
            for row in table.rows:
                for cell in row.cells:
                    # Crie um novo parágrafo para a célula
                    paragraph = cell.paragraphs[0]

                    # Antes de adicionar o texto, adicione a imagem
                    imagem_logo = "Imagens/LOGO.png"  # Substitua pelo caminho da sua imagem
                    run = paragraph.add_run()
                    run.add_picture(imagem_logo, width=Cm(largura_polegadas_logo), height=Cm(altura_polegadas_logo))  # Ajuste as dimensões da imagem

                    # Alinhe o parágrafo ao meio
                    paragraph.alignment = 1  # 1 representa o alinhamento ao meio

                    # Adicione o texto formatado 
                    paragraph.add_run(f"Nome: {nome}\nData: {datas_incrementadas[-1]}\nBenefício: {refeicao}")
                    
                    paragraph = cell.paragraphs[0]

                    # Adicione a imagem da assinatura digital
                    assinatura_path = "Imagens/ASSINATURA.png"  # Substitua pelo caminho correto da imagem
                    paragraph.add_run().add_picture(assinatura_path, width=Cm(largura_polegadas), height=Cm(altura_polegadas))

                    # Incremente a data para o próximo dia
                    dia, mes, ano = map(int, datas_incrementadas[-1].split('/'))
                    dia += 2
                    if dia > 31:  # Simplesmente incrementa o dia (não lida com meses diferentes)
                        dia = 1
                        mes += 1
                    datas_incrementadas.append(f"{dia:02d}/{mes:02d}/{ano}")

            # Chame a função match_case para salvar o documento com base no benefício
            match_case(refeicao, pasta_tickets_cafe, pasta_tickets_almoco, pasta_tickets_janta, nome)